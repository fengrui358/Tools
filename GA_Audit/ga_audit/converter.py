"""Word to PDF converter with filename processing."""

import os
import re
import subprocess
import shutil
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document

from .config import WORD_DIR

# Try to import docx2pdf, fallback to LibreOffice CLI
try:
    from docx2pdf import convert as docx2pdf_convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

# Find LibreOffice installation
LIBREOFFICE_PATH = None
if os.path.exists("/Applications/LibreOffice.app/Contents/MacOS/soffice"):
    LIBREOFFICE_PATH = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
elif shutil.which("soffice"):
    LIBREOFFICE_PATH = "soffice"
elif shutil.which("libreoffice"):
    LIBREOFFICE_PATH = "libreoffice"


def normalize_filename(filename: str) -> str:
    """
    Normalize filename by removing prefix if it doesn't start with 2024年 or 2025年.

    Examples:
        【已检查】2025年12月感知源后端运维文档汇总.docx -> 2025年12月感知源后端运维文档汇总.docx
        前缀内容2024年3月感知源后端运维文档汇总.docx -> 2024年3月感知源后端运维文档汇总.docx
    """
    # Check if filename starts with 2024年 or 2025年
    if filename.startswith("2024年") or filename.startswith("2025年"):
        return filename

    # Find the position of 2024年 or 2025年 in the filename
    match = re.search(r"(2024年|2025年)", filename)
    if match:
        # Remove everything before the match
        return filename[match.start():]

    # If no year pattern found, return original filename
    return filename


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text content from a Word document."""
    try:
        doc = Document(docx_path)
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text.strip())

        return "\n".join(text_content)
    except Exception as e:
        return f"Error reading document: {e}"


def convert_to_pdf(word_path: Path, output_dir: Path) -> Optional[Path]:
    """
    Convert a Word document to PDF.

    Args:
        word_path: Path to the Word document
        output_dir: Directory to save the PDF file

    Returns:
        Path to the converted PDF file, or None if conversion failed
    """
    try:
        # Normalize filename
        original_name = word_path.name
        normalized_name = normalize_filename(original_name)
        pdf_name = normalized_name.replace(".docx", ".pdf")
        pdf_path = output_dir / pdf_name

        # Try docx2pdf first
        if HAS_DOCX2PDF:
            docx2pdf_convert(str(word_path), str(pdf_path))
        # Fallback to LibreOffice CLI
        elif LIBREOFFICE_PATH:
            # Use subprocess to call LibreOffice
            result = subprocess.run(
                [
                    LIBREOFFICE_PATH,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(output_dir),
                    str(word_path)
                ],
                capture_output=True,
                text=True
            )
            # LibreOffice might output a different filename
            # Check for the expected PDF file
            if not pdf_path.exists():
                # Check if there's a PDF with the original filename
                original_pdf = output_dir / word_path.name.replace(".docx", ".pdf")
                if original_pdf.exists():
                    # Rename it to the normalized name
                    original_pdf.rename(pdf_path)
        else:
            print(f"Warning: Neither docx2pdf nor LibreOffice found. Cannot convert {word_path.name}")
            return None

        if pdf_path.exists():
            return pdf_path
        return None
    except Exception as e:
        print(f"Error converting {word_path.name}: {e}")
        return None


def create_pdf_output_dir() -> Path:
    """Create a timestamped PDF output directory."""
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    pdf_dir = Path.cwd() / f"pdf_{timestamp}"
    pdf_dir.mkdir(exist_ok=True)
    return pdf_dir


def get_word_files() -> list[Path]:
    """Get all Word documents from the Word directory."""
    if not WORD_DIR.exists():
        return []

    word_files = list(WORD_DIR.glob("*.docx")) + list(WORD_DIR.glob("*.doc"))
    return sorted(word_files)


def convert_all_word_to_pdf() -> dict[str, Path | str]:
    """
    Convert all Word documents to PDF.

    Returns:
        Dictionary with conversion results
    """
    word_files = get_word_files()

    if not word_files:
        return {"status": "error", "message": "No Word files found in Word directory"}

    if not HAS_DOCX2PDF and not LIBREOFFICE_PATH:
        return {
            "status": "error",
            "message": "Neither docx2pdf nor LibreOffice found. Please install LibreOffice:\n"
                     "  brew install --cask libreoffice\n\n"
                     "PDF conversion is not available, but content audit will still work."
        }

    output_dir = create_pdf_output_dir()
    results = {
        "status": "success",
        "output_dir": output_dir,
        "converted": [],
        "failed": [],
        "renamed": [],
        "conversion_method": "docx2pdf" if HAS_DOCX2PDF else "LibreOffice CLI"
    }

    for word_file in word_files:
        original_name = word_file.name
        normalized_name = normalize_filename(original_name)

        # Track if filename was normalized
        if original_name != normalized_name:
            results["renamed"].append({
                "original": original_name,
                "normalized": normalized_name
            })

        pdf_path = convert_to_pdf(word_file, output_dir)

        if pdf_path:
            results["converted"].append({
                "word_file": original_name,
                "pdf_file": pdf_path.name
            })
        else:
            results["failed"].append(original_name)

    return results
